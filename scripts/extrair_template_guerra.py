# -*- coding: utf-8 -*-
"""Gera o template correto do parecer Guerra Advogados: pega o cabecalho/rodape
(logo GUERRA ADVOGADOS real) do parecer Thomson/Dominio VF_ML.docx e monta um
DOCX limpo, so com header/footer + esqueleto de secoes, para servir de modelo
ao MCP Compliance (layouts_listar -> parecer_guerra).
"""
import copy
from docx import Document
from docx.oxml.ns import qn

FONTE = r"G:\Meu Drive\Projetos de Execução\Consultorias\Guerra\03_Thompson - Dominio Sistemas\NEW 2026\05_PARECER REUNIAO 01062026\TR - PARECER 01062026 VF_ML.docx"
SAIDA = r"G:\Meu Drive\Projetos de Execução\Consultorias\Guerra\00_Modelos\PARECER TÉCNICO - MODELO GUERRA ADVOGADOS.docx"

doc = Document(FONTE)

# limpa o corpo, preservando section properties (header/footer/margens)
body = doc.element.body
sectPr = body.find(qn("w:sectPr"))
for child in list(body):
    if child is not sectPr:
        body.remove(child)


def p_titulo(txt, bold=True, center=True, size=13):
    par = doc.add_paragraph()
    r = par.add_run(txt)
    r.bold = bold
    if size:
        from docx.shared import Pt
        r.font.size = Pt(size)
    if center:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return par


def p_secao(txt):
    par = doc.add_paragraph()
    r = par.add_run(txt)
    r.bold = True
    return par


def p_corpo(txt="[Texto do parecer]"):
    from docx.shared import Cm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    par = doc.add_paragraph()
    par.add_run(txt)
    pf = par.paragraph_format
    pf.line_spacing = Pt(18)
    pf.space_after = Pt(4)
    pf.first_line_indent = Cm(2.0)
    par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return par


p_titulo("PARECER JURÍDICO - TRIBUTÁRIO E FISCAL")

tbl = doc.add_table(rows=4, cols=2)
rotulos = ["CONSULENTE", "DATA", "ASSUNTO", "LEGISLAÇÃO"]
for i, rot in enumerate(rotulos):
    tbl.rows[i].cells[0].text = rot
    tbl.rows[i].cells[0].paragraphs[0].runs[0].bold = True
    tbl.rows[i].cells[1].text = "[preencher]"
doc.add_paragraph()

p_titulo("EMENTA: [SÍNTESE DA MATÉRIA EM CAIXA ALTA]", size=None, center=False)

for numero, nome in [
    ("I.", "RELATÓRIO"),
    ("II.", "FUNDAMENTAÇÃO"),
    ("III.", "CONCLUSÃO"),
]:
    p_secao("%s %s" % (numero, nome))
    p_corpo()

p_titulo("É o parecer, s.m.j.", bold=False, size=None)
p_titulo("Fortaleza, [DATA POR EXTENSO].", bold=False, size=None)

doc.add_paragraph()
for nome, cargo, reg in [
    ("Prof. Fellipe Guerra", "Contador e Advogado Tributarista", "CRC/CE nº 21.074 | OAB/CE nº 49.759"),
    ("Prof. Marcos Lima", "Contador e Cientista de Dados", "CRC/CE nº 23.224"),
]:
    p_titulo("________________________________________", bold=False, size=None)
    p_titulo(nome, size=None)
    p_titulo(cargo, bold=False, size=None)
    p_titulo(reg, bold=False, size=None)

doc.save(SAIDA)
print("OK:", SAIDA)
